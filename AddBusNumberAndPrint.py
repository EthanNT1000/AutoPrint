import os
import docx
import csv
import shutil
import subprocess
import time

def printing(path):
    printer = subprocess.Popen(["powershell", "Start-Process", "-FilePath", path, "-verb", "print"])
    printer.wait()
    #os.startfile(path,"print")
    #time.sleep(5)
def change_paragraphs(file,row):
    for para in file.paragraphs:
        if "EngineNumber" in para.text:
            inline = para.runs
            for i in range(len(inline)):
                if 'EngineNumber' in inline[i].text:
                    text = inline[i].text.replace('EngineNumber', row["EngineNumber"])
                    inline[i].text = text

        if "BodyNumber" in para.text:
            inline = para.runs
            for i in range(len(inline)):
                if "BodyNumber" in inline[i].text:
                    text = inline[i].text.replace("BodyNumber", row["BodyNumber"])
                    inline[i].text = text


with open("number.csv", "r") as csvf:
    reader = csv.DictReader(csvf)
    for row in reader:
        file = docx.Document('01-05.docx')
        # change word
        change_paragraphs(file,row)
        for table in file.tables:
            for cell in table._cells:
                change_paragraphs(cell,row)

        # save files
        os.chdir("AutoPrintData")
        if not os.path.isdir(row['BodyNumber']):
            os.mkdir(row['BodyNumber'])
        os.chdir(row['BodyNumber'])
        file.save('01-05_{}.docx'.format(row['BodyNumber']))

        # copy files
        files = os.listdir("../../files")
        for f in files:
            shutil.copy("../../files/{}".format(f), os.getcwd())

        # print files
        files = os.listdir(os.getcwd())
        rootPath = os.getcwd()
        os.chdir("../..")
        for file in files:
            path = os.path.join(rootPath, file)
            printing(path)


